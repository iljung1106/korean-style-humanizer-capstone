#!/usr/bin/env python3
"""Build a Stage06B Human/AI/Rewrite report from rewrite generations.

The script intentionally keeps the Stage05F report contract:
human_control / ai_source / rewrite_output distributions, Cliff's delta based
filtering, violin plots, DOCX summary, and a side-by-side HTML sample viewer.
"""

from __future__ import annotations

import argparse
import csv
import html
import json
import math
import re
import sys
from collections import defaultdict
from pathlib import Path
from typing import Any


SCRIPT = Path(__file__).resolve()
ROOT = SCRIPT.parents[2]
PHASE_EVAL_ROOT = ROOT / "pipeline_v2" / "eval" / "phase_eval"
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
if str(PHASE_EVAL_ROOT) not in sys.path:
    sys.path.insert(0, str(PHASE_EVAL_ROOT))

from metrics_gui_compatible import cliffs_delta, finite_float, summarize_numeric  # noqa: E402
from score_and_report import build_analysis_rows, numeric_metric_names, score_rows  # noqa: E402


EXCLUDED_METRICS = {
    "raw_chars",
    "analysis_chars",
    "char_len",
    "sentence_count",
    "morph_count",
    "sampled_offset",
    "translationese_raw",
    "content_top_10_coverage",
    "content_gini_frequency",
    "content_repeat_occurrence_rate",
    "comma_sentence_rate",
    "parenthesis_pair_per_1k_chars",
}

EXCLUDED_PREFIXES = ("composite_", "contract_")

SIMILE_PATTERNS = (
    re.compile(r"마치"),
    re.compile(r"처럼"),
    re.compile(r"같이"),
    re.compile(r"듯이"),
    re.compile(r"마냥"),
    re.compile(r"[가-힣A-Za-z0-9]+같은"),
    re.compile(r"[가-힣A-Za-z0-9]+듯한"),
)

GROUP_ORDER = ["human_control", "ai_source", "rewrite_output"]
GROUP_LABELS = {
    "human_control": "Human",
    "ai_source": "AI source",
    "rewrite_output": "Rewrite",
}
GROUP_COLORS = {
    "human_control": "#2563eb",
    "ai_source": "#dc2626",
    "rewrite_output": "#16a34a",
}

FOCUS_METRICS = (
    "comma_per_1k_chars",
    "sentence_initial_token_repeat_rate",
    "parenthesis_pair_per_1k_chars",
    "modifier_repeat_burst_mass",
    "modifier_repetition_mass",
    "simile_marker_per_1k_chars",
    "content_modifier_repeat_occurrence_rate",
    "simile_sentence_rate",
    "pos_3gram_repeat_rate",
    "anti_slop_density",
    "pos_5gram_repeat_rate",
    "pos_4gram_diversity",
    "sentence_length_iqr_ratio",
    "sentence_length_cv",
    "sentence_final_token_repeat_rate",
)

KOREAN_NAMES = {
    "comma_per_1k_chars": "쉼표 밀도",
    "sentence_initial_token_repeat_rate": "문장 시작 토큰 반복률",
    "parenthesis_pair_per_1k_chars": "괄호쌍 밀도",
    "modifier_repeat_burst_mass": "근접 수식어 반복 질량",
    "modifier_repetition_mass": "수식어 반복 질량",
    "simile_marker_per_1k_chars": "직유 표지 밀도",
    "content_modifier_repeat_occurrence_rate": "내용어 중 수식어 반복률",
    "simile_sentence_rate": "직유 표지 문장 비율",
    "pos_3gram_repeat_rate": "POS 3-gram 반복률",
    "anti_slop_density": "AI 상투 표현 밀도",
    "pos_5gram_repeat_rate": "POS 5-gram 반복률",
    "pos_4gram_diversity": "POS 4-gram 다양도",
    "sentence_length_iqr_ratio": "문장 길이 IQR 비율",
    "sentence_length_cv": "문장 길이 변동계수",
    "sentence_final_token_repeat_rate": "문장 마지막 토큰 반복률",
}


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    with path.open(encoding="utf-8") as handle:
        for line in handle:
            line = line.strip()
            if line:
                rows.append(json.loads(line))
    return rows


def write_jsonl(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")


def write_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    keys: list[str] = []
    seen: set[str] = set()
    for row in rows:
        for key in row:
            if key not in seen:
                keys.append(key)
                seen.add(key)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=keys, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def safe_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", value).strip("_") or "metric"


def percentile(values: list[float], q: float) -> float:
    values = sorted(value for value in values if math.isfinite(value))
    if not values:
        return float("nan")
    if len(values) == 1:
        return values[0]
    pos = (len(values) - 1) * q
    lo = math.floor(pos)
    hi = math.ceil(pos)
    if lo == hi:
        return values[lo]
    return values[lo] * (hi - pos) + values[hi] * (pos - lo)


def mean(values: list[float]) -> float:
    values = [value for value in values if math.isfinite(value)]
    return sum(values) / len(values) if values else float("nan")


def std(values: list[float]) -> float:
    values = [value for value in values if math.isfinite(value)]
    if len(values) < 2:
        return float("nan")
    m = mean(values)
    return math.sqrt(sum((value - m) ** 2 for value in values) / (len(values) - 1))


def simile_metrics(text: str) -> dict[str, float | int]:
    text = str(text or "")
    chars = max(1, len(text))
    sentences = [part for part in re.split(r"(?<=[.!?。！？다요죠까네음함임])\s+", text) if part.strip()]
    sentence_count = max(1, len(sentences))
    marker_count = sum(len(pattern.findall(text)) for pattern in SIMILE_PATTERNS)
    sentence_hits = sum(1 for sentence in sentences if any(pattern.search(sentence) for pattern in SIMILE_PATTERNS))
    return {
        "simile_marker_count": marker_count,
        "simile_marker_per_1k_chars": 1000.0 * marker_count / chars,
        "simile_sentence_rate": sentence_hits / sentence_count,
    }


def values_by_group(rows: list[dict[str, Any]], metric: str) -> dict[str, list[float]]:
    out: dict[str, list[float]] = {}
    for group in GROUP_ORDER:
        values = [
            finite_float(row.get(metric))
            for row in rows
            if row.get("eval_task") == "rewrite" and row.get("group") == group
        ]
        out[group] = [value for value in values if math.isfinite(value)]
    return out


def effect_label(delta: float) -> str:
    value = abs(delta)
    if value >= 0.474:
        return "large"
    if value >= 0.33:
        return "medium"
    if value >= 0.147:
        return "small"
    return "negligible"


def build_metric_rows(scored_rows: list[dict[str, Any]], metrics: list[str]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for metric in metrics:
        grouped = values_by_group(scored_rows, metric)
        human = grouped["human_control"]
        ai = grouped["ai_source"]
        rewrite = grouped["rewrite_output"]
        if len(human) < 2 or len(ai) < 2 or len(rewrite) < 2:
            continue
        human_mean = mean(human)
        ai_mean = mean(ai)
        rewrite_mean = mean(rewrite)
        ai_gap = abs(ai_mean - human_mean)
        rewrite_gap = abs(rewrite_mean - human_mean)
        gap_closure = (ai_gap - rewrite_gap) / ai_gap if ai_gap > 1e-12 else float("nan")
        delta_ai_human = cliffs_delta(ai, human)
        delta_rewrite_human = cliffs_delta(rewrite, human)
        delta_rewrite_ai = cliffs_delta(rewrite, ai)
        out.append(
            {
                "metric": metric,
                "metric_ko": KOREAN_NAMES.get(metric, metric),
                "human_n": len(human),
                "ai_n": len(ai),
                "rewrite_n": len(rewrite),
                "human_mean": human_mean,
                "ai_mean": ai_mean,
                "rewrite_mean": rewrite_mean,
                "human_std": std(human),
                "ai_std": std(ai),
                "rewrite_std": std(rewrite),
                "human_q25": percentile(human, 0.25),
                "ai_q25": percentile(ai, 0.25),
                "rewrite_q25": percentile(rewrite, 0.25),
                "human_q50": percentile(human, 0.50),
                "ai_q50": percentile(ai, 0.50),
                "rewrite_q50": percentile(rewrite, 0.50),
                "human_q75": percentile(human, 0.75),
                "ai_q75": percentile(ai, 0.75),
                "rewrite_q75": percentile(rewrite, 0.75),
                "human_ai_cliffs_delta": delta_ai_human,
                "rewrite_human_cliffs_delta": delta_rewrite_human,
                "rewrite_ai_cliffs_delta": delta_rewrite_ai,
                "human_ai_effect": effect_label(delta_ai_human),
                "gap_closure": gap_closure,
                "improved": bool(gap_closure > 0),
                "abs_human_ai_delta": abs(delta_ai_human),
                "abs_rewrite_human_delta": abs(delta_rewrite_human),
                "abs_rewrite_ai_delta": abs(delta_rewrite_ai),
            }
        )
    out.sort(key=lambda row: (row["improved"], row["gap_closure"], row["abs_human_ai_delta"]), reverse=True)
    return out


def metric_is_reportable(metric: str) -> bool:
    if metric not in FOCUS_METRICS:
        return False
    if metric in EXCLUDED_METRICS:
        return False
    if any(metric.startswith(prefix) for prefix in EXCLUDED_PREFIXES):
        return False
    if metric.endswith("_count") and metric not in {"simile_marker_count"}:
        return False
    return True


def plot_metric(rows: list[dict[str, Any]], metric: str, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"rewrite_{safe_name(metric)}.png"
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import seaborn as sns

        flat: list[dict[str, Any]] = []
        for group, values in values_by_group(rows, metric).items():
            flat.extend({"group": GROUP_LABELS[group], "value": value} for value in values)
        fig, ax = plt.subplots(figsize=(9, 5.2))
        sns.violinplot(
            data=flat,
            x="group",
            y="value",
            order=[GROUP_LABELS[group] for group in GROUP_ORDER],
            palette=[GROUP_COLORS[group] for group in GROUP_ORDER],
            inner="quartile",
            cut=0,
            linewidth=1.25,
            saturation=0.95,
            ax=ax,
        )
        sns.stripplot(
            data=flat,
            x="group",
            y="value",
            order=[GROUP_LABELS[group] for group in GROUP_ORDER],
            hue="group",
            palette={GROUP_LABELS[group]: GROUP_COLORS[group] for group in GROUP_ORDER},
            size=2.2,
            alpha=0.65,
            jitter=0.22,
            dodge=False,
            legend=False,
            ax=ax,
        )
        ax.set_title(f"rewrite: {metric}")
        ax.set_xlabel("")
        ax.set_ylabel(metric)
        ax.grid(axis="y", alpha=0.25)
        for tick, group in zip(ax.get_xticklabels(), GROUP_ORDER):
            tick.set_color(GROUP_COLORS[group])
            tick.set_fontweight("bold")
        fig.tight_layout()
        fig.savefig(path, dpi=160)
        plt.close(fig)
        return path
    except Exception:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        plt.close("all")
        data = [values_by_group(rows, metric)[group] for group in GROUP_ORDER]
        fig, ax = plt.subplots(figsize=(9, 5.2))
        ax.violinplot(data, showmeans=True, showmedians=True)
        ax.boxplot(data, widths=0.15, showfliers=False)
        ax.set_xticks(range(1, len(GROUP_ORDER) + 1))
        ax.set_xticklabels([GROUP_LABELS[group] for group in GROUP_ORDER])
        ax.set_title(f"rewrite: {metric}")
        ax.grid(axis="y", alpha=0.25)
        fig.tight_layout()
        fig.savefig(path, dpi=160)
        plt.close(fig)
        return path


def fmt(value: Any, ndigits: int = 4) -> str:
    value = finite_float(value)
    if not math.isfinite(value):
        return "NA"
    return f"{value:.{ndigits}f}"


def metric_explanation(metric: str) -> str:
    explanations = {
        "comma_per_1k_chars": "본문 1,000자당 쉼표 개수입니다. 쉼표 과사용과 번역투식 호흡을 보는 보조 지표입니다.",
        "sentence_initial_token_repeat_rate": "문장 첫 어휘가 반복되는 비율입니다. 문장 시작 패턴의 단조로움을 봅니다.",
        "sentence_final_token_repeat_rate": "문장 마지막 어휘가 반복되는 비율입니다. 종결부 반복과 문장 끝 리듬의 단조로움을 봅니다.",
        "sentence_length_cv": "문장 길이 표준편차를 평균으로 나눈 값입니다. 문장 호흡의 변동성을 봅니다.",
        "sentence_length_iqr_ratio": "문장 길이 25~75% 범위의 폭을 중앙값으로 나눈 값입니다. 극단값보다 중간 분포의 호흡 차이를 봅니다.",
        "modifier_repetition_mass": "수식어가 반복 출현한 양을 전체 수식어량으로 나눈 값입니다.",
        "modifier_repeat_burst_mass": "같은 수식어가 가까운 거리에서 재등장할수록 큰 가중치를 주는 반복 지표입니다.",
        "content_modifier_repeat_occurrence_rate": "내용어 중 수식 계열 표현이 반복 출현하는 비율입니다.",
        "simile_marker_per_1k_chars": "`마치`, `처럼`, `같이`, `듯이`, `마냥`, `-같은`, `-듯한` 등 직유 표지의 1,000자당 빈도입니다.",
        "simile_sentence_rate": "직유 표지를 포함한 문장 비율입니다.",
        "anti_slop_density": "AI 상투 표현 lexicon hit를 길이로 정규화한 밀도입니다.",
        "pos_3gram_repeat_rate": "Kiwi 품사열에서 3개 연속 품사 패턴이 반복되는 비율입니다.",
        "pos_5gram_repeat_rate": "Kiwi 품사열에서 5개 연속 품사 패턴이 반복되는 비율입니다.",
        "pos_4gram_diversity": "Kiwi 품사열에서 서로 다른 4-gram 패턴 비율입니다. 높을수록 품사 구조가 다양합니다.",
    }
    return explanations.get(metric, "GUI 호환 문체 지표입니다. Human/AI/Rewrite 분포 비교를 통해 인간 표본 쪽으로 이동했는지 봅니다.")


def write_docx(path: Path, metric_rows: list[dict[str, Any]], plot_dir: Path, *, title: str) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(f"{title} 보고서", 0)
    doc.add_paragraph(
        "AI 원문 100개를 지정된 merged 모델로 rewrite한 뒤, 같은 수의 Human control, AI source, Rewrite output 분포를 비교했습니다. "
        "표시 지표는 Human과 AI 사이의 Cliff's delta가 medium 이상이고, 길이/중복/composite 제외 기준을 통과한 항목입니다."
    )
    doc.add_paragraph(
        "gap_closure는 평균 기준으로 AI source와 Human 사이의 거리가 Rewrite에서 얼마나 줄었는지를 뜻합니다. "
        "양수면 평균상 개선, 음수면 악화입니다."
    )
    doc.add_heading("요약", 1)
    improved = [row for row in metric_rows if row["improved"]]
    not_improved = [row for row in metric_rows if not row["improved"]]
    doc.add_paragraph(f"개선 지표: {len(improved)}개 / 미개선 또는 악화 지표: {len(not_improved)}개")

    for row in metric_rows:
        metric = str(row["metric"])
        doc.add_page_break()
        doc.add_heading(f"{row['metric_ko']} ({metric})", 1)
        plot_path = plot_dir / f"rewrite_{safe_name(metric)}.png"
        if plot_path.exists():
            doc.add_picture(str(plot_path), width=Inches(6.4))
        doc.add_paragraph(metric_explanation(metric))
        doc.add_paragraph(
            f"해석: gap_closure={fmt(row['gap_closure'])}, "
            f"Human-AI Cliff's delta={fmt(row['human_ai_cliffs_delta'])} ({row['human_ai_effect']}), "
            f"Rewrite-Human Cliff's delta={fmt(row['rewrite_human_cliffs_delta'])}. "
            + ("Rewrite 평균은 AI source보다 Human 평균에 가까워졌습니다." if row["improved"] else "Rewrite 평균은 AI source보다 Human 평균에 가까워지지 않았습니다.")
        )
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["Group", "Mean", "Std", "Q25", "Median", "Q75"]
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header
        for group, label in [("human", "Human"), ("ai", "AI source"), ("rewrite", "Rewrite")]:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = fmt(row[f"{group}_mean"])
            cells[2].text = fmt(row[f"{group}_std"])
            cells[3].text = fmt(row[f"{group}_q25"])
            cells[4].text = fmt(row[f"{group}_q50"])
            cells[5].text = fmt(row[f"{group}_q75"])
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_html_report(path: Path, metric_rows: list[dict[str, Any]], plot_dir: Path, *, title: str) -> None:
    rows_html = []
    for row in metric_rows:
        metric = html.escape(str(row["metric"]))
        plot = plot_dir / f"rewrite_{safe_name(str(row['metric']))}.png"
        rel = plot.relative_to(path.parent) if plot.exists() else plot
        rows_html.append(
            f"<section><h2>{html.escape(str(row['metric_ko']))} <code>{metric}</code></h2>"
            f"<p>gap_closure={fmt(row['gap_closure'])}, Human-AI delta={fmt(row['human_ai_cliffs_delta'])}, "
            f"Rewrite-Human delta={fmt(row['rewrite_human_cliffs_delta'])}</p>"
            f"<img src='{html.escape(rel.as_posix())}'></section>"
        )
    path.write_text(
        f"""<!doctype html><meta charset="utf-8"><title>{html.escape(title)}</title>
<style>
body{{font-family:system-ui,'Malgun Gothic',sans-serif;margin:24px;background:#f8fafc;color:#111827}}
section{{background:white;border:1px solid #d8dee9;border-radius:8px;padding:16px;margin:18px 0}}
img{{max-width:100%;border:1px solid #e5e7eb;background:white}}
code{{font-size:.9em}}
table{{border-collapse:collapse}}td,th{{border:1px solid #ddd;padding:4px 8px}}
</style>
<h1>{html.escape(title)}</h1>
"""
        + "\n".join(rows_html),
        encoding="utf-8",
    )


def result_body(text: str) -> str:
    text = str(text or "")
    match = re.search(r"<result>(.*?)</result>", text, flags=re.S)
    if match:
        return match.group(1).strip()
    return text.strip()


def write_sample_viewer(path: Path, eval_dir: Path, generation_dir: Path, *, title: str) -> None:
    prompts = {row["id"]: row for row in read_jsonl(eval_dir / "rewrite_prompts.jsonl")}
    human_controls_path = eval_dir / "human_controls.jsonl"
    human_controls = read_jsonl(human_controls_path) if human_controls_path.exists() else []
    generations = read_jsonl(generation_dir / "rewrite_generations.jsonl")
    rows: list[str] = []
    enriched: list[dict[str, Any]] = []
    for index, gen in enumerate(generations, start=1):
        prompt = prompts.get(str(gen.get("id")), {})
        source = str(prompt.get("source_text") or gen.get("source_text") or "")
        reference = str(prompt.get("reference_text") or gen.get("reference_text") or "")
        human_control = ""
        if human_controls:
            control = human_controls[(index - 1) % len(human_controls)]
            human_control = str(control.get("text") or control.get("source_text") or control.get("reference_text") or "")
        comparison_text = reference or human_control
        comparison_label = "Human reference" if reference else "Human control (unpaired)"
        rewrite = result_body(str(gen.get("generated_text") or ""))
        enriched.append(
            {
                "id": gen.get("id", ""),
                "source_text": source,
                "reference_text": reference,
                "human_control_text": human_control,
                "rewrite_text": rewrite,
                "prompt": gen.get("prompt") or prompt.get("prompt"),
                "hit_result_close": gen.get("hit_result_close"),
                "generated_tokens": gen.get("generated_tokens"),
            }
        )
        rows.append(
            "<section>"
            f"<h2>{index:03d}. {html.escape(str(gen.get('id','')))}</h2>"
            "<div class='grid'>"
            f"<article><h3>AI source</h3><pre>{html.escape(source)}</pre></article>"
            f"<article><h3>Rewrite</h3><pre>{html.escape(rewrite)}</pre></article>"
            f"<article><h3>{html.escape(comparison_label)}</h3><pre>{html.escape(comparison_text)}</pre></article>"
            "</div></section>"
        )
    write_jsonl(path.with_suffix(".jsonl"), enriched)
    path.write_text(
        f"""<!doctype html><meta charset="utf-8"><title>{html.escape(title)} Samples</title>
<style>
body{{font-family:system-ui,'Malgun Gothic',sans-serif;margin:18px;background:#f3f4f6;color:#111827}}
section{{background:white;border:1px solid #d1d5db;border-radius:8px;margin:16px 0;padding:14px}}
.grid{{display:grid;grid-template-columns:1fr 1fr 1fr;gap:12px}}
article{{border:1px solid #e5e7eb;border-radius:6px;padding:10px;background:#fff}}
h2{{font-size:16px}}h3{{font-size:14px;margin-top:0}}
pre{{white-space:pre-wrap;word-break:break-word;font-family:'Malgun Gothic',system-ui,sans-serif;font-size:14px;line-height:1.55}}
</style>
<h1>{html.escape(title)} Samples</h1>
"""
        + "\n".join(rows),
        encoding="utf-8",
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--eval-dir", required=True)
    parser.add_argument("--generation-dir", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--prefix", default="stage06b_rewrite100")
    parser.add_argument("--title", default="Stage06B Rewrite 100 Human / AI / Rewrite 문체 지표")
    parser.add_argument("--min-effect", choices=["all", "small", "medium", "large"], default="medium")
    parser.add_argument("--require-result-tags", action=argparse.BooleanOptionalAction, default=True)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    raw_dir = out_dir / "metrics"
    plot_dir = out_dir / "plots"
    page_plot_dir = out_dir / "metric_page_plots"
    sample_dir = out_dir / "samples"
    sample_dir.mkdir(parents=True, exist_ok=True)

    score_args = argparse.Namespace(
        eval_dir=args.eval_dir,
        generation_dir=args.generation_dir,
        output_dir=str(raw_dir),
        ai_source_controls="",
        human_controls="",
        rewrite_generations="",
        generate_generations="",
        require_result_tags=args.require_result_tags,
        anti_slop_lexicon=str(ROOT / "data" / "processed" / "anti_slop_lexicon.json"),
        translationese_model=str(ROOT / "models" / "translationese_svm" / "svm_detector.joblib"),
    )
    analysis_rows = build_analysis_rows(score_args)
    scored_rows = score_rows(analysis_rows, score_args)
    text_by_key = {
        (row.get("eval_task"), row.get("group"), row.get("id")): str(row.get("analysis_text") or "")
        for row in analysis_rows
    }
    for row in scored_rows:
        row.update(simile_metrics(text_by_key.get((row.get("eval_task"), row.get("group"), row.get("id")), "")))
    scored_rows = [row for row in scored_rows if row.get("eval_task") == "rewrite"]

    metrics = [metric for metric in numeric_metric_names(scored_rows) if metric_is_reportable(metric)]
    for metric in ("simile_marker_per_1k_chars", "simile_sentence_rate"):
        if metric not in metrics:
            metrics.append(metric)
    metric_rows_all = build_metric_rows(scored_rows, metrics)
    if args.min_effect == "all":
        metric_rows = metric_rows_all
    else:
        min_threshold = {"small": 0.147, "medium": 0.33, "large": 0.474}[args.min_effect]
        metric_rows = [row for row in metric_rows_all if abs(float(row["human_ai_cliffs_delta"])) >= min_threshold]
    metric_rows.sort(key=lambda row: (row["improved"], row["gap_closure"], row["abs_human_ai_delta"]), reverse=True)

    write_csv(raw_dir / "raw_metrics_by_sample.csv", scored_rows)
    prefix = safe_name(str(args.prefix))
    write_csv(out_dir / f"{prefix}_metrics_ranked_all.csv", metric_rows_all)
    write_csv(out_dir / f"{prefix}_style_curated_filtered_metrics.csv", metric_rows)
    write_csv(out_dir / f"{prefix}_style_curated_filtered_improved.csv", [row for row in metric_rows if row["improved"]])
    write_csv(out_dir / f"{prefix}_style_curated_filtered_not_improved.csv", [row for row in metric_rows if not row["improved"]])

    for row in metric_rows:
        src = plot_metric(scored_rows, str(row["metric"]), plot_dir)
        dst = page_plot_dir / src.name
        dst.parent.mkdir(parents=True, exist_ok=True)
        dst.write_bytes(src.read_bytes())

    write_docx(out_dir / f"{prefix}_style_curated_filtered_metric_pages.docx", metric_rows, page_plot_dir, title=str(args.title))
    write_html_report(out_dir / f"{prefix}_style_curated_filtered_report.html", metric_rows, page_plot_dir, title=str(args.title))
    write_sample_viewer(sample_dir / f"{prefix}_side_by_side.html", Path(args.eval_dir), Path(args.generation_dir), title=str(args.title))
    (out_dir / "summary.json").write_text(
        json.dumps(
            {
                "scored_rows": len(scored_rows),
                "metrics_all": len(metric_rows_all),
                "metrics_filtered": len(metric_rows),
                "improved": sum(1 for row in metric_rows if row["improved"]),
                "not_improved": sum(1 for row in metric_rows if not row["improved"]),
                "output_dir": str(out_dir),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    print((out_dir / "summary.json").read_text(encoding="utf-8"))


if __name__ == "__main__":
    main()
