#!/usr/bin/env python3
"""Build the revised Stage05F Human/AI/Rewrite metric report.

The report intentionally follows the existing per-metric DOCX layout, but adds
the missing experimental context, paired rewrite analysis, effect sizes, and
method details requested after review.
"""

from __future__ import annotations

import json
import math
import random
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTDIR = ROOT / "outputs" / "local_reports" / "stage05f_revised_metric_report"
OUTDIR.mkdir(parents=True, exist_ok=True)

TEMPLATE = ROOT / "outputs" / "local_reports" / "인간_AI_Rewrite_지표_설명.docx"
OUTDOC = ROOT / "outputs" / "local_reports" / "인간_AI_Rewrite_지표_설명_수정본.docx"
WORKDOC = OUTDIR / "인간_AI_Rewrite_지표_설명_수정본.docx"

RANKED_PATH = (
    ROOT
    / "outputs"
    / "local_reports"
    / "stage05f_human_ai_rewrite_report_style_curated_filtered"
    / "stage05f_style_curated_filtered_metrics_ranked.csv"
)
RAW_PATH = (
    ROOT
    / "outputs"
    / "local_reports"
    / "stage05f_human_ai_rewrite_report_len3300"
    / "stage05f_len3300_raw_metrics_with_simile.csv"
)
PLOT_DIR = (
    ROOT
    / "outputs"
    / "local_reports"
    / "stage05f_human_ai_rewrite_report_style_curated_filtered"
    / "metric_page_plots"
)
LEXICON_PATH = ROOT / "data" / "processed" / "anti_slop_lexicon.json"


KOREAN = {
    "comma_per_1k_chars": "쉼표 밀도",
    "sentence_initial_token_repeat_rate": "문장 시작 토큰 반복률",
    "parenthesis_pair_per_1k_chars": "괄호쌍 밀도",
    "modifier_repeat_burst_mass": "근접 수식어 반복 질량",
    "modifier_repetition_mass": "수식어 반복 질량",
    "simile_marker_per_1k_chars": "직유 표지 밀도",
    "content_modifier_repeat_occurrence_rate": "내용어 중 수식어 반복률",
    "simile_sentence_rate": "직유 표지 문장 비율",
    "pos_3gram_repeat_rate": "POS 3-gram 반복률",
    "anti_slop_density": "AI 상투 표현 밀도",
    "pos_5gram_repeat_rate": "POS 5-gram 반복률",
    "pos_4gram_diversity": "POS 4-gram 다양도",
    "sentence_length_iqr_ratio": "문장 길이 IQR 비율",
    "sentence_length_cv": "문장 길이 변동계수",
    "sentence_final_token_repeat_rate": "문장 마지막 토큰 반복률",
}


def fmt(value: float | int | None, ndigits: int = 4) -> str:
    if value is None:
        return "NA"
    try:
        value = float(value)
    except Exception:
        return "NA"
    if math.isnan(value) or math.isinf(value):
        return "NA"
    return f"{value:.{ndigits}f}"


def quantile(series: pd.Series, q: float) -> float:
    return float(pd.Series(series).quantile(q))


def cliffs_delta(a: pd.Series, b: pd.Series) -> float:
    aa = [float(x) for x in a.dropna()]
    bb = [float(x) for x in b.dropna()]
    if not aa or not bb:
        return float("nan")
    gt = 0
    lt = 0
    for x in aa:
        for y in bb:
            if x > y:
                gt += 1
            elif x < y:
                lt += 1
    return (gt - lt) / (len(aa) * len(bb))


def effect_label(delta: float) -> str:
    value = abs(float(delta))
    if value >= 0.474:
        return "large"
    if value >= 0.33:
        return "medium"
    if value >= 0.147:
        return "small"
    return "negligible"


def metric_direction(metric: str) -> str:
    if metric == "pos_4gram_diversity":
        return "높을수록 인간 기준에 가까운 경향"
    if metric == "parenthesis_pair_per_1k_chars":
        return "인간 기준 평균에 가까울수록 좋음"
    return "낮을수록 인간 기준에 가까운 경향"


def metric_method(metric: str) -> tuple[str, str]:
    methods = {
        "comma_per_1k_chars": (
            "문자 단위 지표입니다. 분석 대상 본문에서 ASCII 쉼표 `,`와 전각 쉼표 `，`를 직접 셉니다.",
            "comma_count / char_len * 1000. 1,000자당 쉼표 개수입니다.",
        ),
        "sentence_initial_token_repeat_rate": (
            "Kiwi 문장 분리 후 각 문장에서 첫 번째 비문장부호 표면 토큰을 추출합니다.",
            "두 번 이상 등장한 문장 시작 토큰의 출현량 합 / 전체 문장 시작 토큰 수입니다.",
        ),
        "parenthesis_pair_per_1k_chars": (
            "괄호성 기호의 열린/닫힌 쌍 사용을 문자 단위로 셉니다. 평균이 꼬리값에 민감하므로 median과 q75를 함께 봐야 합니다.",
            "parenthesis_pair_count / char_len * 1000입니다.",
        ),
        "modifier_repeat_burst_mass": (
            "Kiwi 형태소 분석에서 MAG, MAJ, MM 계열을 수식어로 보고 표면형/품사 조합을 만듭니다.",
            "같은 수식어가 다시 나올 때 거리 d에 대해 exp(-d/20)을 더한 뒤 수식어 토큰 수로 나눕니다.",
        ),
        "modifier_repetition_mass": (
            "수식어 표면형/품사 조합의 전체 반복량을 봅니다. 단순 빈도보다 반복된 항목의 총량을 봅니다.",
            "수식어별 출현 수에서 첫 출현을 제외한 반복분 합 / 전체 수식어 토큰 수입니다.",
        ),
        "simile_marker_per_1k_chars": (
            "직유 후보 표지를 regex로 셉니다. 패턴은 `마치`, `처럼`, `같이`, `듯이`, `마냥`, `-같은`, `-듯한`입니다.",
            "simile_marker_count / char_len * 1000입니다. false positive 가능성이 있어 보조 지표로 봅니다.",
        ),
        "content_modifier_repeat_occurrence_rate": (
            "내용어 그룹 중 수식어 계열에 해당하는 lemma/품사 항목의 반복 출현률입니다.",
            "두 번 이상 등장한 content modifier 항목의 출현량 합 / content modifier 토큰 수입니다.",
        ),
        "simile_sentence_rate": (
            "Kiwi로 나눈 각 문장에 직유 후보 표지가 하나 이상 있는지 봅니다.",
            "직유 표지를 포함한 문장 수 / 전체 문장 수입니다.",
        ),
        "pos_3gram_repeat_rate": (
            "Kiwi 형태소 품사열에서 문장부호를 제외하고 3개 연속 POS 묶음을 만듭니다.",
            "두 번 이상 등장한 POS 3-gram의 출현량 합 / 전체 POS 3-gram 수입니다.",
        ),
        "anti_slop_density": (
            "AI 소설에서 인간 표본보다 높은 빈도로 나타난 n-gram/skip n-gram lexicon을 사용합니다.",
            "lexicon feature hit count에 feature weight를 곱해 더한 뒤, 전체 feature count로 나누고 1000을 곱합니다.",
        ),
        "pos_5gram_repeat_rate": (
            "Kiwi 품사열에서 5개 연속 POS 묶음을 만듭니다.",
            "두 번 이상 등장한 POS 5-gram의 출현량 합 / 전체 POS 5-gram 수입니다.",
        ),
        "pos_4gram_diversity": (
            "Kiwi 품사열에서 4개 연속 POS 묶음을 만들고 서로 다른 패턴 수를 봅니다.",
            "distinct POS 4-gram 수 / 전체 POS 4-gram 수입니다.",
        ),
        "sentence_length_iqr_ratio": (
            "Kiwi 문장 분리 후 문장별 문자 길이를 계산합니다.",
            "(q75 문장 길이 - q25 문장 길이) / median 문장 길이입니다.",
        ),
        "sentence_length_cv": (
            "Kiwi 문장 분리 후 문장별 문자 길이를 계산합니다.",
            "문장 길이 표준편차 / 평균 문장 길이입니다.",
        ),
        "sentence_final_token_repeat_rate": (
            "Kiwi 문장 분리 후 각 문장에서 마지막 비문장부호 표면 토큰을 추출합니다. 엄밀히 말해 종결 어미 자체가 아니라 마지막 표면 토큰 반복입니다.",
            "두 번 이상 등장한 마지막 토큰의 출현량 합 / 전체 문장 마지막 토큰 수입니다.",
        ),
    }
    return methods.get(metric, ("지표별 준비 방식 설명이 아직 등록되지 않았습니다.", "계산식 설명이 아직 등록되지 않았습니다."))


def metric_interpretation(metric: str, gap_closure: float) -> str:
    specific = {
        "anti_slop_density": "핵심 지표입니다. Rewrite는 AI보다 낮아졌지만 Human과의 거리가 아직 큽니다. 다음 학습에서는 이 축을 계속 밀어야 합니다.",
        "sentence_length_iqr_ratio": "악화 지표입니다. Rewrite가 문장 길이 분포를 인간보다 더 좁히거나 다른 방향으로 움직였을 가능성이 있습니다.",
        "sentence_length_cv": "악화 지표입니다. 문장 호흡이 충분히 다양해졌다고 보기 어렵습니다.",
        "sentence_final_token_repeat_rate": "가장 중요한 악화 축입니다. 문장 끝 표현 반복을 직접 줄이는 프롬프트와 reward가 필요합니다.",
        "pos_4gram_diversity": "개선폭은 작습니다. 단독 판단보다는 POS 반복 계열의 일부로 봐야 합니다.",
        "parenthesis_pair_per_1k_chars": "median이 0인 그룹이 많아 평균 기반 해석이 불안정합니다. 보조 지표로만 봐야 합니다.",
    }
    if metric in specific:
        return specific[metric]
    if gap_closure > 0.3:
        return "Rewrite 평균은 AI 원문보다 인간 평균에 가까워졌습니다. 다만 paired 개선율과 Rewrite vs AI 효과크기를 함께 봐야 합니다."
    if gap_closure > 0:
        return "Rewrite 평균은 개선 방향이지만 개선폭은 제한적입니다. 관련 지표군 안에서 보조적으로 추적하는 편이 안전합니다."
    return "Rewrite 평균은 AI보다 인간 평균에서 더 멀어졌습니다. 후속 prompt/reward에서 별도 제어가 필요합니다."


def build_stats() -> pd.DataFrame:
    ranked = pd.read_csv(RANKED_PATH)
    raw = pd.read_csv(RAW_PATH)
    rng = random.Random(20260604)
    rows = []
    for metric in ranked["metric"].tolist():
        human = raw.loc[raw.group == "human_control", metric].dropna()
        ai = raw.loc[raw.group == "ai_source", metric].dropna()
        rewrite = raw.loc[raw.group == "rewrite_output", metric].dropna()
        gap = abs(ai.mean() - human.mean())
        closure = (gap - abs(rewrite.mean() - human.mean())) / gap if gap else float("nan")
        boots = []
        for _ in range(300):
            h_mean = human.sample(len(human), replace=True, random_state=rng.randint(0, 10**9)).mean()
            a_mean = ai.sample(len(ai), replace=True, random_state=rng.randint(0, 10**9)).mean()
            w_mean = rewrite.sample(len(rewrite), replace=True, random_state=rng.randint(0, 10**9)).mean()
            b_gap = abs(a_mean - h_mean)
            if b_gap:
                boots.append((b_gap - abs(w_mean - h_mean)) / b_gap)

        ai_sub = raw.loc[raw.group == "ai_source", ["chunk_id", metric]].dropna().drop_duplicates("chunk_id")
        rw_sub = raw.loc[raw.group == "rewrite_output", ["chunk_id", metric]].dropna().drop_duplicates("chunk_id")
        paired = ai_sub.merge(rw_sub, on="chunk_id", suffixes=("_source", "_rewrite"))
        human_mean = float(human.mean())
        if len(paired):
            before = (paired[f"{metric}_source"] - human_mean).abs()
            after = (paired[f"{metric}_rewrite"] - human_mean).abs()
            improved = after < before
            paired_closure = ((before - after) / before.replace(0, pd.NA)).dropna()
            paired_improved_rate = float(improved.mean())
            paired_mean_delta = float((paired[f"{metric}_rewrite"] - paired[f"{metric}_source"]).mean())
            paired_gap_closure_median = float(paired_closure.median()) if len(paired_closure) else float("nan")
        else:
            paired_improved_rate = float("nan")
            paired_mean_delta = float("nan")
            paired_gap_closure_median = float("nan")

        rows.append(
            {
                "metric": metric,
                "metric_ko": KOREAN.get(metric, metric),
                "direction": metric_direction(metric),
                "human_mean": human.mean(),
                "ai_mean": ai.mean(),
                "rewrite_mean": rewrite.mean(),
                "human_median": human.median(),
                "ai_median": ai.median(),
                "rewrite_median": rewrite.median(),
                "human_q25": quantile(human, 0.25),
                "ai_q25": quantile(ai, 0.25),
                "rewrite_q25": quantile(rewrite, 0.25),
                "human_q75": quantile(human, 0.75),
                "ai_q75": quantile(ai, 0.75),
                "rewrite_q75": quantile(rewrite, 0.75),
                "human_std": human.std(),
                "ai_std": ai.std(),
                "rewrite_std": rewrite.std(),
                "ai_vs_human_delta": cliffs_delta(ai, human),
                "rewrite_vs_ai_delta": cliffs_delta(rewrite, ai),
                "rewrite_vs_human_delta": cliffs_delta(rewrite, human),
                "gap_closure": closure,
                "gap_ci_low": quantile(pd.Series(boots), 0.025) if boots else float("nan"),
                "gap_ci_high": quantile(pd.Series(boots), 0.975) if boots else float("nan"),
                "paired_n": len(paired),
                "paired_improved_rate": paired_improved_rate,
                "paired_mean_delta": paired_mean_delta,
                "paired_gap_closure_median": paired_gap_closure_median,
                "status": "개선" if closure > 0.05 else ("악화" if closure < -0.05 else "변화 작음"),
            }
        )
    stats = pd.DataFrame(rows).sort_values("gap_closure", ascending=False)
    stats.to_csv(OUTDIR / "stage05f_revised_metric_stats.csv", index=False, encoding="utf-8-sig")
    return stats


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def fix_margins(doc: Document) -> None:
    # The source template contains float-like margin values in some environments.
    # Assigning proper Length objects before table creation avoids python-docx
    # parsing failures when it calculates block width.
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)


def add_table(doc: Document, headers: list[str], data: list[list[str]], font_size: int = 8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for row in data:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def add_small_note(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)


def build_doc(stats: pd.DataFrame) -> None:
    shutil.copyfile(TEMPLATE, WORKDOC)
    doc = Document(WORKDOC)
    fix_margins(doc)
    clear_document_body(doc)

    lexicon = json.loads(LEXICON_PATH.read_text(encoding="utf-8"))
    terms = lexicon.get("terms", [])
    kind_counts = pd.Series([term.get("kind") for term in terms]).value_counts().to_dict()
    term_examples = ", ".join(str(term.get("term", "")) for term in terms[:12])
    source_summary = lexicon.get("source_summary", {})

    doc.add_heading("Stage05F Human / AI / Rewrite 문체 지표 상세 보고서 - 수정본", 0)
    doc.add_paragraph(
        "목적: Stage05F rewrite 결과가 AI 원문보다 인간 소설 분포에 가까워졌는지, 어떤 지표는 개선/악화됐는지, "
        "그리고 다음 GRPO reward 설계에 어떤 지표를 사용할지 판단하기 위한 보고서입니다."
    )
    doc.add_paragraph(
        "이 수정본은 기존 지표별 페이지 양식을 유지하되, 데이터셋 정의, Rewrite 생성 방식, paired 분석, "
        "통계적 불확실성, 형태소/문장 분리 기준, anti-slop lexicon 설명을 보강했습니다."
    )
    doc.add_paragraph(
        "중요 제한: 모든 지표는 문체적 proxy입니다. 수치가 인간 평균에 가까워졌다는 사실이 곧 서사 품질 전체의 개선을 의미하지는 않습니다."
    )

    doc.add_heading("1. 실험 개요와 데이터 정의", 1)
    add_table(
        doc,
        ["그룹", "n", "정의", "비고"],
        [
            ["Human", "300", "로컬 인간 웹소설 raw chunk에서 추출한 대조군", "Stage05F rewrite 입력과 직접 paired되지는 않음"],
            ["AI", "300", "로컬 AI 소설 raw chunk에서 추출한 원문 대조군", "Rewrite 100개는 이 중 chunk_id로 매칭 가능한 일부에서 생성됨"],
            ["Rewrite", "100", "Stage05F merged 모델이 AI 원문 일부를 rewrite한 결과", "prompt와 result contract를 사용해 생성된 산출물"],
        ],
    )
    doc.add_paragraph("Rewrite 100개는 `chunk_id` 기준으로 AI source와 매칭됩니다. 각 지표 페이지의 paired 분석은 같은 원문이 rewrite 후 인간 평균에 더 가까워졌는지 계산합니다.")
    doc.add_paragraph("Stage05F는 pipeline_v2의 rewrite GRPO 이후 merge된 모델 결과를 의미합니다. 이 보고서에서는 평가 대상 rewrite 모델 버전으로 정의합니다.")

    doc.add_heading("2. 전처리와 분석 파이프라인", 1)
    doc.add_paragraph("분석 텍스트: result 태그가 있는 생성물은 result 내부 본문을 우선 사용하고, Human/AI raw chunk는 분석용 정제 본문을 사용합니다.")
    doc.add_paragraph("문장 분리: `kiwipiepy.Kiwi.split_into_sents`를 우선 사용합니다. Kiwi가 없을 경우에만 구두점/줄바꿈 기반 regex fallback을 사용합니다.")
    doc.add_paragraph("형태소 분석: `kiwipiepy.Kiwi.tokenize` 결과의 표면형과 품사 태그를 사용합니다. POS n-gram, 수식어 계열, 문장 시작/마지막 토큰 지표는 이 결과에 민감합니다.")
    doc.add_paragraph("gap closure 공식: `(abs(AI_mean - Human_mean) - abs(Rewrite_mean - Human_mean)) / abs(AI_mean - Human_mean)`. 1은 인간 평균 접근, 0은 개선 없음, 음수는 악화입니다.")
    doc.add_paragraph("Cliff's delta: 분포 효과크기입니다. 이 보고서에는 AI vs Human뿐 아니라 Rewrite vs AI, Rewrite vs Human도 함께 기록합니다.")
    doc.add_paragraph("통계적 불확실성: gap closure는 bootstrap 300회로 95% 구간을 추정했습니다.")

    doc.add_heading("3. Anti-slop lexicon 요약", 1)
    doc.add_paragraph(f"사전 파일: {LEXICON_PATH}")
    doc.add_paragraph(f"lexicon term 수: {len(terms)}개. kind별 개수: {kind_counts}.")
    doc.add_paragraph(
        f"구축 corpus 요약: human_docs={source_summary.get('human_docs')}, ai_docs={source_summary.get('ai_docs')}, "
        f"human_tokens={source_summary.get('human_tokens')}, ai_tokens={source_summary.get('ai_tokens')}."
    )
    doc.add_paragraph(f"상위 term 예시: {term_examples}")
    doc.add_paragraph("계산 시에는 한글 2~12자 토큰을 뽑고 unigram, bigram, skip-bigram, skip-trigram feature를 만든 뒤 lexicon hit에 가중치를 곱해 density를 계산합니다.")

    doc.add_heading("4. 전체 요약", 1)
    add_table(
        doc,
        ["지표", "상태", "Human", "AI", "Rewrite", "gap closure", "95% CI", "paired 개선율", "Rewrite vs AI δ"],
        [
            [
                row.metric_ko,
                row.status,
                fmt(row.human_mean),
                fmt(row.ai_mean),
                fmt(row.rewrite_mean),
                fmt(row.gap_closure, 3),
                f"{fmt(row.gap_ci_low, 3)}~{fmt(row.gap_ci_high, 3)}",
                fmt(row.paired_improved_rate, 3),
                fmt(row.rewrite_vs_ai_delta, 3),
            ]
            for row in stats.itertuples()
        ],
        font_size=7,
    )
    doc.add_paragraph("주요 개선: 쉼표 밀도, 문장 시작 반복률, 수식어 반복 burst, 직유 표지 계열은 AI보다 인간 평균에 가까워졌습니다.")
    doc.add_paragraph("주요 미해결: anti-slop density는 낮아졌지만 여전히 Human과 큰 차이가 있고, 문장 길이 IQR/CV 및 문장 마지막 토큰 반복률은 악화됐습니다.")

    for index, row in enumerate(stats.itertuples(), start=1):
        doc.add_page_break()
        metric = row.metric
        doc.add_heading(f"{index:02d}. {row.metric_ko}", 1)
        doc.add_paragraph(
            f"{metric} | {row.status} | gap closure {fmt(row.gap_closure * 100, 1)}% | paired 개선율 {fmt(row.paired_improved_rate * 100, 1)}%"
        )
        plot_candidates = list(PLOT_DIR.glob(f"*_{metric}.png"))
        if plot_candidates:
            doc.add_picture(str(plot_candidates[0]), width=Inches(6.2))
        prep, calc = metric_method(metric)
        doc.add_paragraph("지표별 준비 방식")
        doc.add_paragraph(prep)
        doc.add_paragraph("계산 방식")
        doc.add_paragraph(calc)
        doc.add_paragraph("목표 방향")
        doc.add_paragraph(metric_direction(metric))
        doc.add_paragraph("지표 결과 해설")
        doc.add_paragraph(metric_interpretation(metric, row.gap_closure))
        doc.add_paragraph("이번 결과")
        doc.add_paragraph(
            f"Human mean={fmt(row.human_mean)}, AI mean={fmt(row.ai_mean)}, Rewrite mean={fmt(row.rewrite_mean)}. "
            f"gap closure={fmt(row.gap_closure, 3)}이며 bootstrap 95% 구간은 {fmt(row.gap_ci_low, 3)}~{fmt(row.gap_ci_high, 3)}입니다. "
            f"AI vs Human Cliff's delta={fmt(row.ai_vs_human_delta, 3)}({effect_label(row.ai_vs_human_delta)}), "
            f"Rewrite vs AI delta={fmt(row.rewrite_vs_ai_delta, 3)}({effect_label(row.rewrite_vs_ai_delta)}), "
            f"Rewrite vs Human delta={fmt(row.rewrite_vs_human_delta, 3)}({effect_label(row.rewrite_vs_human_delta)})입니다. "
            f"매칭된 {int(row.paired_n)}개 원문 중 인간 평균에 가까워진 비율은 {fmt(row.paired_improved_rate * 100, 1)}%입니다."
        )
        add_table(
            doc,
            ["group", "mean", "median", "q25", "q75", "std", "n"],
            [
                ["Human", fmt(row.human_mean), fmt(row.human_median), fmt(row.human_q25), fmt(row.human_q75), fmt(row.human_std), "300"],
                ["AI", fmt(row.ai_mean), fmt(row.ai_median), fmt(row.ai_q25), fmt(row.ai_q75), fmt(row.ai_std), "300"],
                ["Rewrite", fmt(row.rewrite_mean), fmt(row.rewrite_median), fmt(row.rewrite_q25), fmt(row.rewrite_q75), fmt(row.rewrite_std), "100"],
            ],
            font_size=8,
        )
        add_small_note(
            doc,
            "주의: 표본 수와 길이 분포 차이, 문장 수가 적은 샘플, 특정 장르/작품 편향은 지표를 흔들 수 있습니다. paired 개선율은 같은 AI 원문이 실제로 개선됐는지 보기 위한 보조 판단입니다.",
        )

    fix_margins(doc)
    doc.save(WORKDOC)
    shutil.copyfile(WORKDOC, OUTDOC)


def main() -> None:
    stats = build_stats()
    build_doc(stats)
    print(
        json.dumps(
            {
                "docx": str(OUTDOC),
                "workdoc": str(WORKDOC),
                "stats_csv": str(OUTDIR / "stage05f_revised_metric_stats.csv"),
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
